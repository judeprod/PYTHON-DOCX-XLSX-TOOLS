"""
Batch DOCX Formatter - Complete Version
Copies formatting from a template document and applies it to multiple DOCX files.
Searches through all subfolders and preserves folder structure.
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os
from pathlib import Path

class DocxFormatter:
    def __init__(self, template_path):
        """Initialize with template document path"""
        self.template = Document(template_path)
        self.styles = self._extract_styles()
        
    def _extract_styles(self):
        """Extract font and paragraph styles from template"""
        styles = {
            'normal': {},
            'heading1': {},
            'heading2': {},
            'heading3': {}
        }
        
        # Extract styles from template paragraphs
        for para in self.template.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                style_name = para.style.name.lower().replace(' ', '')
                
                if 'heading1' in style_name:
                    styles['heading1'] = self._get_run_format(first_run)
                elif 'heading2' in style_name:
                    styles['heading2'] = self._get_run_format(first_run)
                elif 'heading3' in style_name:
                    styles['heading3'] = self._get_run_format(first_run)
                elif not styles['normal']:
                    styles['normal'] = self._get_run_format(first_run)
        
        return styles
    
    def _get_run_format(self, run):
        """Extract formatting properties from a run"""
        return {
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
        }
    
    def _apply_run_format(self, run, format_dict):
        """Apply formatting to a run"""
        if format_dict.get('font_name'):
            run.font.name = format_dict['font_name']
        if format_dict.get('font_size'):
            run.font.size = format_dict['font_size']
        if format_dict.get('bold') is not None:
            run.font.bold = format_dict['bold']
        if format_dict.get('italic') is not None:
            run.font.italic = format_dict['italic']
        if format_dict.get('color'):
            run.font.color.rgb = format_dict['color']
    
    def format_document(self, input_path, output_path, preserve_emphasis=False, 
                        bold_headings=False, bold_first_line=False):
        """Apply template formatting to a document"""
        doc = Document(input_path)
        
        for para in doc.paragraphs:
            style_name = para.style.name.lower().replace(' ', '')
            
            # Determine target style
            if 'heading1' in style_name or 'heading 1' in style_name:
                target_style = self.styles['heading1']
                is_heading = True
            elif 'heading2' in style_name or 'heading 2' in style_name:
                target_style = self.styles['heading2']
                is_heading = True
            elif 'heading3' in style_name or 'heading 3' in style_name:
                target_style = self.styles['heading3']
                is_heading = True
            else:
                target_style = self.styles['normal']
                is_heading = False
            
            # Apply formatting to each run
            for i, run in enumerate(para.runs):
                original_bold = run.font.bold
                original_italic = run.font.italic
                
                # Apply font name and size from template
                if target_style.get('font_name'):
                    run.font.name = target_style['font_name']
                if target_style.get('font_size'):
                    run.font.size = target_style['font_size']
                if target_style.get('color'):
                    run.font.color.rgb = target_style['color']
                
                # Now handle bold/italic explicitly
                if preserve_emphasis:
                    # Keep original bold/italic ONLY if it was explicitly True
                    if original_bold is True:
                        run.font.bold = True
                    elif original_bold is False:
                        run.font.bold = False
                    else:
                        # If None/undefined, use template setting or default to False
                        run.font.bold = target_style.get('bold', False)
                    
                    if original_italic is True:
                        run.font.italic = True
                    elif original_italic is False:
                        run.font.italic = False
                    else:
                        run.font.italic = target_style.get('italic', False)
                else:
                    # Don't preserve - use template only
                    run.font.bold = target_style.get('bold', False)
                    run.font.italic = target_style.get('italic', False)
                
                # Override: Force bold on headings if enabled
                if bold_headings and is_heading:
                    run.font.bold = True
                
                # Override: Bold first run if enabled
                if bold_first_line and i == 0 and not is_heading:
                    run.font.bold = True
        
        doc.save(output_path)
        print(f"✓ Formatted: {os.path.basename(input_path)}")
    
    def batch_format(self, input_folder, output_folder, recursive=True,
                    preserve_emphasis=False, bold_headings=False, bold_first_line=False):
        """Format all DOCX files in a folder"""
        input_path = Path(input_folder)
        output_path = Path(output_folder)
        
        output_path.mkdir(exist_ok=True)
        
        # Find all DOCX files
        if recursive:
            docx_files = list(input_path.rglob("*.docx"))
        else:
            docx_files = list(input_path.glob("*.docx"))
        
        # Filter out temp files and output folder files
        docx_files = [f for f in docx_files 
                     if not f.name.startswith("~$") 
                     and not str(f).startswith(str(output_path))]
        
        if not docx_files:
            print(f"\nNo DOCX files found in {input_folder}")
            return
        
        print(f"\nFound {len(docx_files)} document(s) to format\n")
        
        # Process each file
        for docx_file in docx_files:
            try:
                if recursive:
                    rel_path = docx_file.relative_to(input_path)
                    out_file = output_path / rel_path
                    out_file.parent.mkdir(parents=True, exist_ok=True)
                else:
                    out_file = output_path / docx_file.name
                
                self.format_document(str(docx_file), str(out_file), 
                                   preserve_emphasis, bold_headings, bold_first_line)
            except Exception as e:
                print(f"✗ Error formatting {docx_file.name}: {str(e)}")
        
        print(f"\n✓ All done! Check: {output_folder}")


def main():
    """Main execution function"""
    print("=" * 60)
    print("DOCX Batch Formatter")
    print("=" * 60)
    
    # Set root folder
    root_folder = r"C:\Users\judep\Downloads\SMS FOR EDITING"
    
    if not os.path.exists(root_folder):
        print(f"\nError: Root folder not found: {root_folder}")
        input("\nPress Enter to exit...")
        return
    
    print(f"\nRoot folder: {root_folder}")
    
    # Step 1: Get template file
    print("\n" + "=" * 60)
    print("STEP 1: Select Template File")
    print("=" * 60)
    print("Tip: Drag and drop the file or type the path")
    template_path = input("\nEnter template DOCX path: ").strip()
    
    # Remove quotes
    template_path = template_path.strip('"').strip("'")
    
    # Check if just filename
    if not os.path.isabs(template_path):
        template_path = os.path.join(root_folder, template_path)
    
    if not os.path.exists(template_path):
        print(f"\nError: Template file not found!")
        input("\nPress Enter to exit...")
        return
    
    print(f"✓ Using template: {os.path.basename(template_path)}")
    
    # Step 2: Choose location
    print("\n" + "=" * 60)
    print("STEP 2: Select Input Location")
    print("=" * 60)
    print("1. Format ALL DOCX files (root folder + all subfolders)")
    print("2. Format files in a specific subfolder (+ its subfolders)")
    print("3. Format only root folder (no subfolders)")
    print("4. Enter custom folder path (drag & drop)")
    
    choice = input("\nEnter choice (1-4): ").strip()
    
    recursive = True
    
    if choice == "2":
        # List subfolders
        try:
            subfolders = [f for f in os.listdir(root_folder) 
                         if os.path.isdir(os.path.join(root_folder, f)) 
                         and f != "formatted_output"]
            
            if subfolders:
                print("\nAvailable subfolders:")
                for i, sf in enumerate(subfolders, 1):
                    print(f"  {i}. {sf}")
                
                sf_choice = input("\nEnter number or folder name: ").strip()
                
                try:
                    idx = int(sf_choice) - 1
                    if 0 <= idx < len(subfolders):
                        subfolder = subfolders[idx]
                    else:
                        subfolder = sf_choice
                except ValueError:
                    subfolder = sf_choice
            else:
                subfolder = input("Enter subfolder name: ").strip()
        except Exception:
            subfolder = input("Enter subfolder name: ").strip()
        
        input_folder = os.path.join(root_folder, subfolder)
        
        if not os.path.exists(input_folder):
            print(f"\nError: Folder not found!")
            input("\nPress Enter to exit...")
            return
            
    elif choice == "3":
        input_folder = root_folder
        recursive = False
        
    elif choice == "4":
        print("\nTip: Drag and drop the folder")
        input_folder = input("Enter folder path: ").strip()
        input_folder = input_folder.strip('"').strip("'")
        
        if not os.path.exists(input_folder):
            print(f"\nError: Folder not found!")
            input("\nPress Enter to exit...")
            return
    else:
        input_folder = root_folder
    
    print(f"\n✓ Input folder: {input_folder}")
    
    if recursive:
        print("✓ Searching all subfolders")
    else:
        print("✓ Searching this folder only")
    
    # Step 3: Bold/Italic options
    print("\n" + "=" * 60)
    print("STEP 3: Bold & Italic Settings")
    print("=" * 60)
    print("1. Remove ALL bold/italic - plain text only (Recommended)")
    print("2. Keep original bold/italic from documents")
    print("3. Force bold on all headings only")
    print("4. Custom settings")
    
    emphasis_choice = input("\nEnter choice (1-4): ").strip()
    
    preserve_emphasis = False  # Default to removing bold
    bold_headings = False
    bold_first_line = False
    
    if emphasis_choice == "2":
        preserve_emphasis = True
        print("✓ Will keep original bold/italic from documents")
    elif emphasis_choice == "3":
        preserve_emphasis = False
        bold_headings = True
        print("✓ Will remove bold/italic except headings")
    elif emphasis_choice == "4":
        print("\nCustom settings:")
        keep_orig = input("Keep original bold/italic from docs? (y/n): ").strip().lower()
        preserve_emphasis = (keep_orig == 'y')
        
        force_heads = input("Force bold on all headings? (y/n): ").strip().lower()
        bold_headings = (force_heads == 'y')
        
        first_line = input("Bold first line of paragraphs? (y/n): ").strip().lower()
        bold_first_line = (first_line == 'y')
        
        print(f"✓ Custom settings applied")
    else:
        print("✓ Will remove all bold/italic - plain text")
    
    output_folder = os.path.join(root_folder, "formatted_output")
    
    print("\n" + "=" * 60)
    print("Processing...")
    print("=" * 60)
    
    # Format documents
    try:
        formatter = DocxFormatter(template_path)
        formatter.batch_format(input_folder, output_folder, recursive,
                             preserve_emphasis, bold_headings, bold_first_line)
        
        print("\n" + "=" * 60)
        print("SUCCESS!")
        print("=" * 60)
        input("\nPress Enter to exit...")
        
    except Exception as e:
        print(f"\n✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        input("\nPress Enter to exit...")


if __name__ == "__main__":
    try:
        import docx
    except ImportError:
        print("\n✗ Error: python-docx not installed!")
        print("Run: pip install python-docx")
        input("\nPress Enter to exit...")
        exit(1)
    
    main()