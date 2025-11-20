"""
Batch DOCX Formatter
Copies formatting from a template document and applies it to multiple DOCX files.
Preserves lists, numbering, bullets, and heading styles.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path
import shutil

class DocxFormatter:
    def __init__(self, template_path):
        """Initialize with template document path"""
        self.template = Document(template_path)
        self.styles = self._extract_styles()
        
    def _extract_styles(self):
        """Extract font and paragraph styles from template"""
        styles = {
            'normal': {},
            'heading1': {},
            'heading2': {},
            'heading3': {}
        }
        
        # Extract styles from template paragraphs
        for para in self.template.paragraphs:
            if para.runs:
                first_run = para.runs[0]
                style_name = para.style.name.lower().replace(' ', '')
                
                if 'heading1' in style_name:
                    styles['heading1'] = self._get_run_format(first_run)
                elif 'heading2' in style_name:
                    styles['heading2'] = self._get_run_format(first_run)
                elif 'heading3' in style_name:
                    styles['heading3'] = self._get_run_format(first_run)
                elif not styles['normal']:  # Get first normal paragraph
                    styles['normal'] = self._get_run_format(first_run)
        
        return styles
    
    def _get_run_format(self, run):
        """Extract formatting properties from a run"""
        return {
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
        }
    
    def _apply_run_format(self, run, format_dict):
        """Apply formatting to a run"""
        if format_dict.get('font_name'):
            run.font.name = format_dict['font_name']
        if format_dict.get('font_size'):
            run.font.size = format_dict['font_size']
        if format_dict.get('bold') is not None:
            run.font.bold = format_dict['bold']
        if format_dict.get('italic') is not None:
            run.font.italic = format_dict['italic']
        if format_dict.get('color'):
            run.font.color.rgb = format_dict['color']
    
    def format_document(self, input_path, output_path):
        """Apply template formatting to a document"""
        doc = Document(input_path)
        
        for para in doc.paragraphs:
            # Determine which style to apply
            style_name = para.style.name.lower().replace(' ', '')
            
            if 'heading1' in style_name or 'heading 1' in style_name:
                target_style = self.styles['heading1']
            elif 'heading2' in style_name or 'heading 2' in style_name:
                target_style = self.styles['heading2']
            elif 'heading3' in style_name or 'heading 3' in style_name:
                target_style = self.styles['heading3']
            else:
                target_style = self.styles['normal']
            
            # Apply formatting to all runs in paragraph
            for run in para.runs:
                # Preserve original bold/italic if they exist
                original_bold = run.font.bold
                original_italic = run.font.italic
                
                self._apply_run_format(run, target_style)
                
                # Restore original emphasis if it was set
                if original_bold is not None:
                    run.font.bold = original_bold
                if original_italic is not None:
                    run.font.italic = original_italic
        
        # Save formatted document
        doc.save(output_path)
        print(f"✓ Formatted: {os.path.basename(input_path)} -> {os.path.basename(output_path)}")
    
    def batch_format(self, input_folder, output_folder=None):
        """Format all DOCX files in a folder"""
        input_path = Path(input_folder)
        
        # Create output folder if not specified
        if output_folder is None:
            output_folder = input_path / "formatted_output"
        else:
            output_folder = Path(output_folder)
        
        output_folder.mkdir(exist_ok=True)
        
        # Find all DOCX files
        docx_files = list(input_path.glob("*.docx"))
        
        # Filter out temporary files
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        
        if not docx_files:
            print(f"No DOCX files found in {input_folder}")
            return
        
        print(f"\nFound {len(docx_files)} document(s) to format")
        print(f"Output folder: {output_folder}\n")
        
        # Process each file
        for docx_file in docx_files:
            try:
                output_path = output_folder / docx_file.name
                self.format_document(str(docx_file), str(output_path))
            except Exception as e:
                print(f"✗ Error formatting {docx_file.name}: {str(e)}")
        
        print(f"\n✓ Batch formatting complete! Check: {output_folder}")


def main():
    """Main execution function"""
    print("=" * 60)
    print("DOCX Batch Formatter")
    print("=" * 60)
    
    # Set root folder
    root_folder = r"C:\Users\judep\Downloads\SMS FOR EDITING"
    
    if not os.path.exists(root_folder):
        print(f"Error: Root folder not found: {root_folder}")
        return
    
    print(f"\nRoot folder: {root_folder}")
    
    # Get template file
    print("\n--- Step 1: Select Template File ---")
    template_path = input("Enter template DOCX filename (in root folder) or full path: ").strip()
    
    # Check if it's just a filename or full path
    if not os.path.isabs(template_path):
        template_path = os.path.join(root_folder, template_path)
    
    if not os.path.exists(template_path):
        print(f"Error: Template file not found: {template_path}")
        return
    
    print(f"✓ Using template: {os.path.basename(template_path)}")
    
    # Option to format files in root or subfolder
    print("\n--- Step 2: Select Input Location ---")
    print("1. Format all DOCX files in root folder")
    print("2. Format files in a subfolder")
    choice = input("Enter choice (1 or 2): ").strip()
    
    if choice == "2":
        subfolder = input("Enter subfolder name: ").strip()
        input_folder = os.path.join(root_folder, subfolder)
        if not os.path.exists(input_folder):
            print(f"Error: Subfolder not found: {input_folder}")
            return
    else:
        input_folder = root_folder
    
    print(f"✓ Input folder: {input_folder}")
    
    # Output folder
    output_folder = os.path.join(root_folder, "formatted_output")
    
    # Create formatter and process
    try:
        formatter = DocxFormatter(template_path)
        formatter.batch_format(input_folder, output_folder)
    except Exception as e:
        print(f"\nError: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # Required: pip install python-docx
    try:
        import docx
    except ImportError:
        print("Error: python-docx not installed!")
        print("Please run: pip install python-docx")
        exit(1)
    
    main()