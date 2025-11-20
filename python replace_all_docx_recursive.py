import os
from docx import Document
import sys
import re
import subprocess

print(sys.executable)

# SETTINGS – edit these before running
root_folder = r"C:\Users\judep\Downloads\FORMS EDITING\UNLOCKED"
find_text = "VEM FORMS REMOVED FOR MANUAL REVISION"
replace_text = "Mention intentionally removed."

# --- Do not edit below this line ---
count_files = 0
count_replaced = 0
count_converted = 0

# Compile a regex pattern for case-insensitive find
pattern = re.compile(re.escape(find_text), re.IGNORECASE)

def replace_text_in_paragraph(paragraph, pattern, replace_text):
    """Simple text replacement in paragraph."""
    if pattern.search(paragraph.text):
        # Get the full text
        full_text = paragraph.text
        # Replace it
        new_text = pattern.sub(replace_text, full_text)
        # Clear the paragraph
        paragraph.clear()
        # Add the new text
        paragraph.add_run(new_text)
        return True
    return False

def convert_doc_to_docx(doc_path):
    """Convert .doc file to .docx using LibreOffice or MS Word if available."""
    try:
        docx_path = doc_path.replace(".doc", ".docx")
        subprocess.run([
            "soffice", "--headless", "--convert-to", "docx", "--outdir",
            os.path.dirname(doc_path), doc_path
        ], capture_output=True, timeout=30)
        
        if os.path.exists(docx_path):
            return docx_path
    except Exception:
        pass
    
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        docx_path = doc_path.replace(".doc", ".docx")
        word.Documents.Open(os.path.abspath(doc_path))
        word.ActiveDocument.SaveAs(os.path.abspath(docx_path), FileFormat=12)
        word.ActiveDocument.Close()
        word.Quit()
        return docx_path
    except Exception:
        pass
    
    return None

for foldername, subfolders, filenames in os.walk(root_folder):
    for filename in filenames:
        if filename.endswith(".docx") or filename.endswith(".doc"):
            file_path = os.path.join(foldername, filename)
            
            if filename.endswith(".doc"):
                print(f"🔄 Converting {filename} to .docx...")
                converted_path = convert_doc_to_docx(file_path)
                if converted_path:
                    file_path = converted_path
                    count_converted += 1
                    print(f"✅ Converted: {converted_path}")
                else:
                    print(f"⚠️ Could not convert {filename}. Skipping.")
                    continue
            
            try:
                doc = Document(file_path)
            except Exception as e:
                print(f"⚠️ Skipped {filename} (error reading file: {e})")
                continue

            replaced_in_file = False

            # Replace in paragraphs
            for paragraph in doc.paragraphs:
                if replace_text_in_paragraph(paragraph, pattern, replace_text):
                    replaced_in_file = True

            # Replace in headers and footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        if replace_text_in_paragraph(para, pattern, replace_text):
                            replaced_in_file = True
                    
                    for table in section.header.tables:
                        try:
                            if table._cells:
                                for cell in table._cells:
                                    for para in cell.paragraphs:
                                        if replace_text_in_paragraph(para, pattern, replace_text):
                                            replaced_in_file = True
                        except (IndexError, AttributeError):
                            continue
                
                if section.footer:
                    for para in section.footer.paragraphs:
                        if replace_text_in_paragraph(para, pattern, replace_text):
                            replaced_in_file = True
                    
                    for table in section.footer.tables:
                        try:
                            if table._cells:
                                for cell in table._cells:
                                    for para in cell.paragraphs:
                                        if replace_text_in_paragraph(para, pattern, replace_text):
                                            replaced_in_file = True
                        except (IndexError, AttributeError):
                            continue

            # Replace in tables
            for table in doc.tables:
                try:
                    if not table._cells:
                        continue
                    for cell in table._cells:
                        for para in cell.paragraphs:
                            if replace_text_in_paragraph(para, pattern, replace_text):
                                replaced_in_file = True
                except (IndexError, AttributeError) as e:
                    print(f"⚠️ Skipped malformed table in {filename}: {e}")
                    continue

            if replaced_in_file:
                doc.save(file_path)
                count_replaced += 1
                print(f"✅ Modified: {file_path}")
            else:
                print(f"❌ No change: {file_path}")

            count_files += 1

print(f"✅ Processed {count_files} Word files (including subfolders).")
print(f"🔄 Converted {count_converted} .doc files to .docx.")
print(f"📝 Updated {count_replaced} files containing '{find_text}'.")
print("Done!")