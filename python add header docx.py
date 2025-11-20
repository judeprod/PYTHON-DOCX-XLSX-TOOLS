import os
import re
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# ---------- SETTINGS ----------
root_folder = r"C:\Users\judep\Downloads\FORMS EDITING\BIGLILLY"
header_rev = "Rev 0 / 2025-11-17"
# ------------------------------

# Patterns to remove from footer
footer_patterns = [
    r"Page\s+\d+\s+of\s+\d+",
    r"Issue\s+Number:\s*\d+",
    r"Revision\s+Number:\s*\d+"
]

def add_field(run, field_type):
    """Adds a Word field like PAGE or NUMPAGES."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.text = field_type
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def clean_footer(section):
    """Remove Page X of Y, Issue Number, Revision Number from footer."""
    footer = section.footer

    for paragraph in footer.paragraphs:
        text = paragraph.text
        
        # Check each pattern
        for pattern in footer_patterns:
            if re.search(pattern, text, flags=re.IGNORECASE):
                print("   Removing footer text →", repr(text.strip()))
                paragraph.clear()  # Clear entire paragraph
                break

def process_file(full_path):
    if "~$" in full_path:
        print("Skipping temporary file:", full_path)
        return

    try:
        doc = Document(full_path)
    except PackageNotFoundError:
        print("Skipping (corrupted or password protected):", full_path)
        return
    except Exception:
        print("Skipping (likely password protected):", full_path)
        return

    section = doc.sections[0]

    # ------ CLEAN FOOTER ------
    clean_footer(section)

    # ------ SET HEADER ------
    header = section.header
    header.paragraphs[0].clear()
    p = header.paragraphs[0]
    p.alignment = 0  # left
    

    # SSP line
    run1 = p.add_run("SSP\n")
    run1.bold = True
    run1.font.name = "Tahoma"
    run1.font.size = Pt(8)

    # Revision line
    run2 = p.add_run(f"{header_rev}\n")
    run2.bold = True
    run2.font.name = "Tahoma"
    run2.font.size = Pt(8)

    # Page X of Y
    run3 = p.add_run("Page ")
    run3.bold = True
    run3.font.name = "Tahoma"
    run3.font.size = Pt(8)

    add_field(run3, "PAGE")

    run4 = p.add_run(" of ")
    run4.bold = True
    run4.font.name = "Tahoma"
    run4.font.size = Pt(8)

    add_field(run4, "NUMPAGES")

    doc.save(full_path)
    print("Updated:", full_path)

# -------- MAIN LOOP --------
for dirpath, dirs, files in os.walk(root_folder):
    for file in files:
        if file.lower().endswith(".docx"):
            full_path = os.path.join(dirpath, file)
            process_file(full_path)

print("\nDone!")
